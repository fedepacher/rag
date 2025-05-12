import logging
import os
from io import BytesIO
from typing import List
import PyPDF2
from langchain.text_splitter import TokenTextSplitter
from docx import Document as Docx


class Document:
    def __init__(self, filename: List[str], file_data: List[bytes]):
        self.filename = filename
        self.file_data = file_data

    def file_type(self, filename: str):
        _, file_extension = os.path.splitext(filename)
        return file_extension.lower()

    def textualize_file(self):
        full_content = ''
        for index in range(len(self.file_data)):
            if ".pdf" == self.file_type(self.filename[index]):
                with BytesIO(self.file_data[index]) as stream:
                    pdf_reader = PyPDF2.PdfReader(stream)
                    text = ""
                    for page in pdf_reader.pages:
                        text += "\n"
                        text += page.extract_text()
                full_content += ''.join(text)
            elif ".docx" == self.file_type(self.filename[index]):
                with BytesIO(self.file_data[index]) as stream:
                    doc = Docx(stream)
                    text = [para.text for para in doc.paragraphs]
                    # Extracting text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    text.append(para.text)
                    text = '\n'.join(text)
                full_content += '\n'.join(text)
            else:
                raise ValueError(f"Unsupported file type: {self.file_type}")
        return full_content

    def chunk_text(self, text, context_length):
        text_splitter = TokenTextSplitter(encoding_name="cl100k_base", chunk_size=context_length, chunk_overlap=200)
        chunked_text = text_splitter.split_text(text)
        return chunked_text

    def get_file_content(self) -> str:
        file_text = self.textualize_file()

        return file_text

    def get_chunked_text(self, context_length):
        merged_text = self.get_file_content()
        chunked_text = self.chunk_text(merged_text, context_length)

        return chunked_text


class DocumentLoader:
    def __init__(self):
        pass

    def load_document(self) -> Document:
        pass


class LocalDocumentLoader(DocumentLoader):
    def __init__(self, base_directory: str):
        super().__init__()
        self.base_directory = base_directory

    def load_document(self) -> Document:
        documents = []
        files = []
        for filename in os.listdir(self.base_directory):
            with open(os.path.join(self.base_directory, filename), 'rb') as f:
                doc = f.read()
            files.append(filename)
            documents.append(doc)
        return Document(files, documents)
